"""
Resume Parser Module
Extracts structured information from resume files (PDF and DOCX)
"""

import os
import re
from PyPDF2 import PdfReader
from docx import Document
from utils import extract_email, extract_phone, normalize_text, extract_name_from_text


class ResumeParser:
    """Parse resume files and extract relevant information"""
    
    def __init__(self):
        """Initialize parser"""
        self.email_pattern = r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}'
        self.phone_pattern = r'(\+\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}'
    
    def parse_resume(self, filepath):
        """
        Parse resume file and extract information
        
        Args:
            filepath (str): Path to resume file
        
        Returns:
            dict: Extracted resume data
        """
        # Check file extension
        file_ext = os.path.splitext(filepath)[1].lower()
        
        if file_ext == '.pdf':
            text = self._extract_text_from_pdf(filepath)
        elif file_ext in ['.docx', '.doc']:
            text = self._extract_text_from_docx(filepath)
        else:
            raise ValueError(f"Unsupported file format: {file_ext}")
        
        # Parse extracted text
        resume_data = {
            'text': text,
            'name': self._extract_name(text),
            'email': self._extract_email(text),
            'phone': self._extract_phone(text),
            'education': self._extract_education(text),
            'experience': self._extract_experience(text),
            'experience_summary': self._extract_experience_summary(text),
            'sections': self._identify_sections(text)
        }
        
        return resume_data
    
    def _extract_text_from_pdf(self, filepath):
        """
        Extract text from PDF file
        
        Args:
            filepath (str): Path to PDF file
        
        Returns:
            str: Extracted text
        """
        try:
            text = []
            with open(filepath, 'rb') as file:
                pdf_reader = PdfReader(file)
                for page in pdf_reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text.append(page_text)
            
            return '\n'.join(text)
        except Exception as e:
            raise Exception(f"Error extracting text from PDF: {str(e)}")
    
    def _extract_text_from_docx(self, filepath):
        """
        Extract text from DOCX file
        
        Args:
            filepath (str): Path to DOCX file
        
        Returns:
            str: Extracted text
        """
        try:
            doc = Document(filepath)
            text = []
            
            # Extract from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text.append(paragraph.text)
            
            # Extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text.append(cell.text)
            
            return '\n'.join(text)
        except Exception as e:
            raise Exception(f"Error extracting text from DOCX: {str(e)}")
    
    def _extract_name(self, text):
        """
        Extract candidate name from resume text
        
        Args:
            text (str): Resume text
        
        Returns:
            str: Extracted name or None
        """
        try:
            name = extract_name_from_text(text)
            return name if name else 'Unknown'
        except:
            return 'Unknown'
    
    def _extract_email(self, text):
        """
        Extract email address
        
        Args:
            text (str): Resume text
        
        Returns:
            str: Email address or None
        """
        try:
            return extract_email(text)
        except:
            return None
    
    def _extract_phone(self, text):
        """
        Extract phone number
        
        Args:
            text (str): Resume text
        
        Returns:
            str: Phone number or None
        """
        try:
            return extract_phone(text)
        except:
            return None
    
    def _extract_education(self, text):
        """
        Extract education information
        
        Args:
            text (str): Resume text
        
        Returns:
            str: Education summary
        """
        education_keywords = [
            'bachelor', 'master', 'phd', 'b.tech', 'm.tech', 'b.sc', 'm.sc',
            'diploma', 'degree', 'bca', 'mca', 'bba', 'mba', 'graduation'
        ]
        
        lines = text.split('\n')
        education_info = []
        
        in_education_section = False
        
        for line in lines:
            line_lower = line.lower()
            
            # Check for education section header
            if any(keyword in line_lower for keyword in ['education', 'academic', 'qualification']):
                in_education_section = True
                continue
            
            # Check for next section
            if in_education_section and any(keyword in line_lower for keyword in ['experience', 'skills', 'project', 'achievements']):
                in_education_section = False
                continue
            
            # Extract education details
            if in_education_section and line.strip():
                education_info.append(line.strip())
            elif any(keyword in line_lower for keyword in education_keywords):
                education_info.append(line.strip())
        
        return ' | '.join(education_info[:3]) if education_info else 'Not Specified'
    
    def _extract_experience(self, text):
        """
        Extract work experience
        
        Args:
            text (str): Resume text
        
        Returns:
            list: List of experience entries
        """
        experience_keywords = [
            'experience', 'employment', 'work', 'position', 'job', 'internship'
        ]
        
        lines = text.split('\n')
        experience_info = []
        
        in_experience_section = False
        
        for line in lines:
            line_lower = line.lower()
            
            # Check for experience section header
            if any(keyword in line_lower for keyword in experience_keywords):
                in_experience_section = True
                continue
            
            # Check for next section
            if in_experience_section and any(keyword in line_lower for keyword in ['education', 'skills', 'project', 'achievements']):
                in_experience_section = False
                continue
            
            # Extract experience details
            if in_experience_section and line.strip():
                experience_info.append(line.strip())
        
        return experience_info[:5] if experience_info else []
    
    def _extract_experience_summary(self, text):
        """
        Generate summary of experience
        
        Args:
            text (str): Resume text
        
        Returns:
            str: Experience summary
        """
        experience_list = self._extract_experience(text)
        
        if experience_list:
            return ' | '.join(experience_list[:2])
        else:
            return 'No experience found'
    
    def _identify_sections(self, text):
        """
        Identify major sections in resume
        
        Args:
            text (str): Resume text
        
        Returns:
            dict: Sections found with their presence
        """
        text_lower = text.lower()
        
        sections = {
            'objective': 'objective' in text_lower or 'summary' in text_lower,
            'education': 'education' in text_lower or 'academic' in text_lower,
            'experience': 'experience' in text_lower or 'employment' in text_lower,
            'skills': 'skills' in text_lower or 'technical' in text_lower,
            'projects': 'project' in text_lower,
            'certifications': 'certification' in text_lower or 'certified' in text_lower,
            'achievements': 'achievement' in text_lower or 'award' in text_lower,
            'languages': 'language' in text_lower,
            'references': 'reference' in text_lower
        }
        
        return sections
